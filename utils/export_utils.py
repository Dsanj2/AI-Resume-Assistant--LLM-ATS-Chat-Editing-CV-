import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import streamlit as st

# =========================
# Helpers
# =========================
def is_header(line: str) -> bool:
    headers = [
        "summary", "experience", "education", "skills",
        "projects", "achievements", "interests / hobbies"
    ]
    return line.strip().lower() in headers

def is_contact_info(line: str) -> bool:
    return any(x in line for x in ["@", "http", "+", "www"])

def clean_line(line: str) -> str:
    return line.strip()

# =========================
# PDF Export
# =========================
def export_pdf(text: str, filename: str = "resume.pdf"):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=40, leftMargin=40,
                            topMargin=40, bottomMargin=40)

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle('Normal', fontSize=10, leading=12)
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Heading1'],
        fontSize=14,
        leading=16,
        textColor=colors.HexColor("#4a07f2"),
        bold=True,
        underline=True,
        spaceAfter=6,
        spaceBefore=6
    )
    contact_style = ParagraphStyle(
        'Contact',
        fontSize=12,
        leading=14,
        textColor=colors.HexColor("#4a07f2"),
        bold=True,
        underline=True,
        spaceAfter=4,
        spaceBefore=4
    )
    name_style = ParagraphStyle(
        'Name',
        fontSize=16,
        leading=18,
        textColor=colors.HexColor("#4a07f2"),
        bold=True,
        underline=True,
        spaceAfter=8,
        spaceBefore=8
    )

    story = []
    lines = text.splitlines()
    for i, line in enumerate(lines):
        clean = clean_line(line)
        if not clean:
            story.append(Spacer(1, 6))
            continue

        # First line assumed to be Name
        if i == 0:
            story.append(Paragraph(clean, name_style))
        elif is_contact_info(clean):
            story.append(Paragraph(clean, contact_style))
        elif is_header(clean):
            story.append(Paragraph(clean, header_style))
        else:
            story.append(Paragraph(clean, normal_style))

    doc.build(story)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Download PDF",
        data=buffer,
        file_name=filename,
        mime="application/pdf",
        use_container_width=True
    )

# =========================
# DOCX Export
# =========================
def export_docx(text: str, filename: str = "resume.docx"):
    buffer = io.BytesIO()
    doc = Document()

    normal_font_size = Pt(10)
    header_font_size = Pt(14)
    contact_font_size = Pt(12)
    name_font_size = Pt(16)
    highlight_color = RGBColor(74, 7, 242)  # Hex #4a07f2

    lines = text.splitlines()
    for i, line in enumerate(lines):
        clean = clean_line(line)
        if not clean:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        run = p.add_run(clean)

        if i == 0:  # First line = Name
            run.bold = True
            run.underline = True
            run.font.size = name_font_size
            run.font.color.rgb = highlight_color
        elif is_contact_info(clean):
            run.bold = True
            run.underline = True
            run.font.size = contact_font_size
            run.font.color.rgb = highlight_color
        elif is_header(clean):
            run.bold = True
            run.underline = True
            run.font.size = header_font_size
            run.font.color.rgb = highlight_color
        else:
            run.font.size = normal_font_size

        # Force Arial font
        run.font.name = 'Arial'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Download DOCX",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
